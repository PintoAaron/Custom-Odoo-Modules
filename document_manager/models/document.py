from odoo import models, fields, api
import base64
from io import BytesIO
from docx import Document
from datetime import datetime

class DocumentManager(models.Model):
    _name = 'document.manager'
    _description = 'Document Manager'

    name = fields.Char(string='Document Name', required=True)
    file = fields.Binary(string='Upload File', attachment=True, required=True)
    file_name = fields.Char(string='File Name')
    file_type = fields.Selection([('pdf', 'PDF'), ('docx', 'DOCX')], string='File Type', required=True)
    last_edit = fields.Datetime(string='Last Edit', readonly=True)
    content = fields.Text(string='Content')

    def action_edit_document(self):
        for record in self:
            if record.file_type == 'docx':
                file_content = base64.b64decode(record.file)
                doc = Document(BytesIO(file_content))
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                record.content = '\n'.join(full_text)
                record.last_edit = fields.Datetime.now()

                
    def action_save_document(self):
        for record in self:
            if record.file_type == 'docx':
                doc = Document()
                for line in record.content.split('\n'):
                    doc.add_paragraph(line)
                file_stream = BytesIO()
                doc.save(file_stream)
                record.file = base64.b64encode(file_stream.getvalue())
    
    
    @api.model
    def create(self, vals):
        vals['last_edit'] = fields.Datetime.now()
        return super(DocumentManager, self).create(vals)

    def write(self, vals):
        vals['last_edit'] = fields.Datetime.now()
        return super(DocumentManager, self).write(vals)